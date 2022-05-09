import sys, os, time, datetime
import docx                                                                                            
from   docx.shared          import RGBColor, Pt                                                      
from   docx.enum.text       import WD_ALIGN_PARAGRAPH
from   colorama             import *
from   art                  import *
################################################################################################################# A  J #####################################################################################################################
###VARS
# path  var
here       = os.getcwd()
# clear var (used print(clear) to clear space)
clear      = "\033[2J\033[1;1f"
# time variables
today      = datetime.date.today().strftime("%d-%m-%Y")
now        = str(datetime.datetime.today().strftime("%H:%M:%S %d-%m-%Y"))

# defining AJ function
def AJ():
       newentry = ""
       print(clear)
       print("\n"*10)
       print("""






                                                                                                                  
                                                                                                         """+Fore.LIGHTRED_EX+"""   ##############      """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""    ##################  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""  #              #     """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+""" #                #    """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""####################   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""####################   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""                     #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""   #                 #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""    #                #  """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""     #              #   """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #   """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""      #            #    """+Style.RESET_ALL+"""
                                                                                                         """+Fore.LIGHTRED_EX+"""#                  #  """+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""#       ############     #"""+Style.RESET_ALL+Style.BRIGHT+"""                                                         """+Fore.LIGHTBLACK_EX+"""["""+Fore.LIGHTRED_EX+"""A"""+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""utomated """+Fore.LIGHTRED_EX+"""J"""+Style.RESET_ALL+Fore.LIGHTBLACK_EX+"""ournal] ("""+Fore.LIGHTRED_EX+"""v7.0"""+Style.RESET_ALL+Fore.LIGHTBLACK_EX+""")"""+Style.RESET_ALL+"""
|-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------|""")
       print("")
       print(now+Style.RESET_ALL)
       newentry = input("")
       if   os.path.exists("Archives\\AJ.docx") == True:
              file              = docx.Document("Archives\\AJ.docx")
              AJ1               = file.add_paragraph().add_run(now)
              AJ2               = file.add_paragraph()
              AJ3               = AJ2.add_run(newentry) 
              AJ1.bold          = True
              AJ3.add_break()
              file.save("Archives\\AJ.docx")
       elif os.path.exists("Archives\\AJ.docx") == False:
              file              = docx.Document()
              AJ0               = file.add_paragraph().add_run("                     AJ")
              AJ0.font.size     = Pt(41)
              AJ0.bold          = True
              AJ0run            = file.add_paragraph().add_run("______________________________________________")
              AJ0run.font.size  = Pt(25)
              AJ0run.bold       = True
              AJ1               = file.add_paragraph().add_run()
              AJ1.add_break()
              AJ1.add_break()
              AJ2               = file.add_paragraph().add_run(now)
              AJ3               = file.add_paragraph().add_run(newentry) 
              AJ2.bold          = True
              AJ3.add_break()
              file.save("Archives\\AJ.docx")
       print(clear)
       print("\n"*20)
       print(Fore.LIGHTGREEN_EX+"")
       tprint("       AJ ENTRY SAVED!", font="cyber")
       if os.path.exists('log.txt') == False:
           log = open('log.txt','w')
           log.write(now + " |  New AJ-Entry.")
           log.close()
       elif os.path.exists('log.txt') == True:
           log = open('log.txt','a')
           log.write("\n" + now + " |  New AJ-Entry.")
           log.close()
       print(""+Style.RESET_ALL)
       time.sleep(0.5)
       

#opening journal and counting the appearances of journal entrys from today in it
def check_journal():
       global jcount
       global jjcount
       if   os.path.exists("Archives\\AJ.docx") == True:
              doc = docx.Document("Archives\\AJ.docx")
              fullText = ""
              for para in doc.paragraphs:
                     fullText = fullText + str(para.text)
              found  = re.findall(str(today), str(fullText))
       elif os.path.exists("Archives\\AJ.docx") == False:
              found = []
       jcount = 0
       for i in found:
           jcount += 1
       jjcount = "Entries"
       if str(jcount) == "1":
              jjcount = "Entry"


AJ()